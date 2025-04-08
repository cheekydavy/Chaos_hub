import os
from flask import Flask, request, render_template, redirect, url_for, send_from_directory, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_socketio import SocketIO, emit, join_room, leave_room
from werkzeug.utils import secure_filename
from datetime import datetime, timezone
from docx import Document
import pandas as pd
import requests
import logging

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-fuckin-key')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///nexus_matrix.db').replace('postgres://', 'postgresql://')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.environ.get('UPLOAD_FOLDER', 'uploads')

db = SQLAlchemy(app)
socketio = SocketIO(app, cors_allowed_origins=['http://localhost:5000', 'http://127.0.0.1:5000'])
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', 'AIzaSyB4CGLzAqNfi2XBRY5DcArjSV7V-_-Irto')  # Replace!
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash-latest:generateContent"

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

upload_folder = app.config['UPLOAD_FOLDER']
if not os.path.exists(upload_folder):
    os.makedirs(upload_folder)

chat_history = []

class Unit(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    notes = db.relationship('File', backref='unit', lazy=True)

class File(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(100), nullable=False)
    type = db.Column(db.String(20), nullable=False)
    upload_date = db.Column(db.String(20), default=lambda: datetime.now(timezone.utc).strftime('%d/%m/%Y'))
    unit_id = db.Column(db.Integer, db.ForeignKey('unit.id'), nullable=True)

class Notice(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    topic = db.Column(db.String(100), nullable=False)
    remark = db.Column(db.String(200), nullable=False)
    posted_date = db.Column(db.String(20), default=lambda: datetime.now(timezone.utc).strftime('%d/%m/%Y'))
    messages = db.relationship('Message', backref='notice', lazy=True)

class Message(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.String(500), nullable=False)
    posted_date = db.Column(db.String(20), default=lambda: datetime.now(timezone.utc).strftime('%d/%m/%Y'))
    notice_id = db.Column(db.Integer, db.ForeignKey('notice.id'), nullable=False)

class Assignment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    topic = db.Column(db.String(100), nullable=False)
    remark = db.Column(db.String(200), nullable=False)
    due_date = db.Column(db.String(20), nullable=False)
    posted_date = db.Column(db.String(20), default=lambda: datetime.now(timezone.utc).strftime('%d/%m/%Y'))

LIBRARY = {
    "Python Programming": {"title": "Python Crash Course", "file": "python_crash_course.pdf"},
    "Informatics": {"title": "Informatics Essentials", "file": "informatics_essentials.pdf"}
}

with app.app_context():
    db.drop_all()
    db.create_all()

def allowed_file(filename):
    return filename.endswith(('.xlsx', '.csv', '.docx', '.pdf'))

def get_ai_response(question):
    headers = {"Content-Type": "application/json"}
    context = "\n".join(chat_history[-4:])
    payload = {
        "contents": [{"parts": [{"text": f"You are a knowledgeable AI assistant for a school. Answer fully and clearly with relevant details. Use this context if relevant:\n{context}\n\nQuestion: {question}"}]}],
        "generationConfig": {"maxOutputTokens": 500}
    }
    logger.debug(f"Sending to Gemini: {payload['contents'][0]['parts'][0]['text']}")
    try:
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers=headers, json=payload)
        response.raise_for_status()
        answer = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
        logger.debug(f"Gemini says: {answer}")
        return answer
    except requests.exceptions.HTTPError as e:
        logger.error(f"Gemini fucked up with status {e.response.status_code}: {e.response.text}")
        return f"Gemini fucked up with status {e.response.status_code}: {e.response.text}"
    except Exception as e:
        logger.error(f"Gemini fucked up: {str(e)}")
        return f"Gemini fucked up: {str(e)}"

def simplify_timetable(df, original_filename, year, semester, course):
    doc = Document()
    doc.add_heading(f'{course} Timetable - Year {year} Semester {semester}', 0)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)
    output_filename = f"simplified_{year}_{semester}_{course}_{original_filename.split('.')[0]}.docx"
    output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
    doc.save(output_path)
    return output_filename

@app.route('/', methods=['GET', 'POST'])
def index():
    units = Unit.query.all()
    class_timetable = File.query.filter_by(type='class_timetable').first()
    exam_timetable = File.query.filter_by(type='exam_timetable').first()
    notices = Notice.query.order_by(Notice.posted_date.desc()).all()
    assignments = Assignment.query.order_by(Assignment.due_date).all()

    if request.method == 'POST':
        if 'note' in request.files:
            file = request.files['note']
            unit_id = request.form['unit_id']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                note = File(filename=filename, type='note', unit_id=unit_id)
                db.session.add(note)
                db.session.commit()
        elif 'class_timetable' in request.files:
            file = request.files['class_timetable']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                if class_timetable:
                    db.session.delete(class_timetable)
                timetable = File(filename=filename, type='class_timetable')
                db.session.add(timetable)
                db.session.commit()
        elif 'exam_timetable' in request.files:
            file = request.files['exam_timetable']
            if file and allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                if exam_timetable:
                    db.session.delete(exam_timetable)
                timetable = File(filename=filename, type='exam_timetable')
                db.session.add(timetable)
                db.session.commit()
        elif 'notice_topic' in request.form:
            topic = request.form['notice_topic']
            remark = request.form['notice_remark']
            notice = Notice(topic=topic, remark=remark)
            db.session.add(notice)
            db.session.commit()
            return redirect(url_for('chat_room', notice_id=notice.id))
        elif 'assignment_topic' in request.form:
            topic = request.form['assignment_topic']
            remark = request.form['assignment_remark']
            due_date = request.form['assignment_due_date']
            if due_date:
                due_date = datetime.strptime(due_date, '%Y-%m-%d').strftime('%d/%m/%Y')
                assignment = Assignment(topic=topic, remark=remark, due_date=due_date)
                db.session.add(assignment)
                db.session.commit()
        return redirect(url_for('index'))

    today = datetime.now(timezone.utc).strftime('%d/%m/%Y')
    for assignment in assignments:
        if datetime.strptime(assignment.due_date, '%d/%m/%Y') < datetime.strptime(today, '%d/%m/%Y'):
            db.session.delete(assignment)
    db.session.commit()
    assignments = Assignment.query.order_by(Assignment.due_date).all()

    return render_template('index.html', units=units, class_timetable=class_timetable,
                          exam_timetable=exam_timetable, notices=notices, assignments=assignments)

@app.route('/chat_room/<int:notice_id>', methods=['GET', 'POST'])
def chat_room(notice_id):
    notice = Notice.query.get_or_404(notice_id)
    if request.method == 'POST':
        content = request.form['message']
        message = Message(content=content, notice_id=notice_id)
        db.session.add(message)
        db.session.commit()
        socketio.emit('message', {'room': str(notice_id), 'msg': f"{message.content} ({message.posted_date})"})
        return redirect(url_for('chat_room', notice_id=notice_id))
    messages = Message.query.filter_by(notice_id=notice_id).order_by(Message.posted_date).all()
    return render_template('chat_room.html', notice=notice, messages=messages)

@app.route('/group_setup', methods=['GET', 'POST'])
def group_setup():
    if request.method == 'POST':
        units = request.form.getlist('units[]')
        for unit_name in units:
            unit = Unit(name=unit_name)
            db.session.add(unit)
        db.session.commit()
        return redirect(url_for('index'))
    return render_template('group_setup.html')

@app.route('/ai_chat', methods=['GET', 'POST'])
def ai_chat():
    if request.method == 'POST':
        question = request.form.get('question')
        file = request.files.get('file')
        response = ""

        logger.debug(f"POST received - Question: {question}, File: {file}")

        if question and not file:
            response = get_ai_response(question)
            if not response:
                response = "Fake AI: I got your question!"
            logger.debug(f"Emitting AI response: {response}")
            socketio.emit('ai_response', {'msg': response})
            chat_history.append(f"User: {question}")
            chat_history.append(f"AI: {response}")

        if question and "book" in question.lower():
            for key, book in LIBRARY.items():
                if key.lower() in question.lower() or book['title'].lower() in question.lower():
                    response = f"Found: {book['title']} - <a href='/uploads/{book['file']}'>Download</a>"
                    logger.debug(f"Emitting library response: {response}")
                    socketio.emit('ai_response', {'msg': response})
                    chat_history.append(f"User: {question}")
                    chat_history.append(f"AI: {response}")
                    break
            else:
                response = "Book not found in the library. Try another title."
                logger.debug(f"Emitting library response: {response}")
                socketio.emit('ai_response', {'msg': response})
                chat_history.append(f"User: {question}")
                chat_history.append(f"AI: {response}")

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            logger.debug(f"File saved: {filepath}")

            if question and "timetable" in question.lower() and "simplify" in question.lower():
                try:
                    df = pd.read_excel(filepath) if filename.endswith('.xlsx') else pd.read_csv(filepath)
                    user_input = question.lower()
                    if "year 2 sem 2" in user_input and "informatics" in user_input:
                        simplified_df = df[(df['Year'] == 2) & (df['Semester'] == 2) & df['Course'].str.contains('Informatics', case=False)]
                        output_file = simplify_timetable(simplified_df, filename, 2, 2, "Informatics")
                        response = f"Simplified timetable ready: <a href='/downloads/{output_file}'>Download</a>"
                        logger.debug(f"Emitting timetable response: {response}")
                        socketio.emit('ai_response', {'msg': response})
                        chat_history.append(f"User: {question}")
                        chat_history.append(f"AI: {response}")
                    else:
                        response = "Please specify 'Year 2 Sem 2 Informatics' for timetable simplification."
                        logger.debug(f"Emitting timetable response: {response}")
                        socketio.emit('ai_response', {'msg': response})
                        chat_history.append(f"User: {question}")
                        chat_history.append(f"AI: {response}")
                except Exception as e:
                    response = f"Error processing file: {str(e)}"
                    logger.error(f"File processing failed: {str(e)}")
                    socketio.emit('ai_response', {'msg': response})
                    chat_history.append(f"User: {question}")
                    chat_history.append(f"AI: {response}")
            else:
                response = "Upload successful. Ask to simplify the timetable if needed."
                logger.debug(f"Emitting upload response: {response}")
                socketio.emit('ai_response', {'msg': response})
                chat_history.append(f"User: {question or 'File upload'}")
                chat_history.append(f"AI: {response}")

        return render_template('ai_chat.html')
    logger.debug("Emitting test message on load")
    socketio.emit('ai_response', {'msg': 'AI Chat Loaded - Ask away!'})
    return render_template('ai_chat.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/downloads/<filename>')
def download_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename, as_attachment=True)

@socketio.on('connect')
def handle_connect():
    logger.debug("Client connected to Socket.IO")
    socketio.emit('ai_response', {'msg': 'AI Chat Loaded - Ask away!'})

@socketio.on('join')
def on_join(data):
    room = request.args.get('notice_id')
    join_room(room)
    emit('joined', {'room': room}, room=room)

@socketio.on('leave')
def on_leave(data):
    room = data['room']
    leave_room(room)

@socketio.on('message')
def handle_message(data):
    room = data['room']
    msg = data['msg']
    emit('message', {'room': room, 'msg': msg}, room=room)

@socketio.on('ai_message')
def handle_ai_message(data):
    question = data['msg']
    logger.debug(f"Received ai_message: {question}")
    response = get_ai_response(question)
    logger.debug(f"Emitting ai_response: {response}")
    emit('ai_response', {'msg': response})
    chat_history.append(f"User: {question}")
    chat_history.append(f"AI: {response}")

if __name__ == '__main__':
    socketio.run(app, host='0.0.0.0', port=5000, debug=True)